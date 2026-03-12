"""
Life Of Research — By Pranay Mahendrakar
GAME-LIKE INTERACTIVE EDITION — Production Release v5.0

A 14-Agent Autonomous AI Research Pipeline that generates full academic papers
using GPT-4o with iterative peer review, debate mode, humanization engine,
and journal-quality evaluation.

Installation:
    pip install PySide6 requests python-docx

Usage:
    python life_of_research_app.py

GitHub: https://github.com/PranayMahendrakar/life-of-research
Author: Pranay Mahendrakar
License: MIT
"""

import sys, re, os, json, math, random, requests
from datetime import datetime
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QTextEdit, QTextBrowser, QPushButton,
    QComboBox, QSpinBox, QDoubleSpinBox, QScrollArea, QFrame,
    QSplitter, QProgressBar, QGroupBox, QFileDialog, QCheckBox,
    QSizePolicy, QTabWidget, QPlainTextEdit, QDialog, QDialogButtonBox
)
from PySide6.QtCore import (Qt, QThread, Signal, QTimer, QRect,
                             QPoint, QPropertyAnimation, QEasingCurve,
                             QRectF, QPointF, QSizeF)
from PySide6.QtGui import (
    QFont, QColor, QPalette, QPainter, QBrush, QPen,
    QLinearGradient, QRadialGradient, QTextCursor,
    QPainterPath, QPolygonF, QTransform
)

# ═══════════════════════════════════════════════════════════════════
#   MARKDOWN → HTML RENDERER
# ═══════════════════════════════════════════════════════════════════

def md_to_html(text: str, accent="#60a5fa") -> str:
    """Convert markdown text to styled HTML for QTextBrowser."""
    lines = text.split("\n")
    html_lines = []
    in_table = False
    in_code  = False
    in_ul    = False
    in_ol    = False

    def close_lists():
        nonlocal in_ul, in_ol
        out = ""
        if in_ul: out += "</ul>"; in_ul = False
        if in_ol: out += "</ol>"; in_ol = False
        return out

    for raw in lines:
        line = raw
        if line.strip().startswith("```"):
            if in_code:
                html_lines.append("</code></pre>"); in_code = False
            else:
                html_lines.append(close_lists())
                html_lines.append(f'<pre style="background:#0a0a18;color:#a8d8a8;'
                                  f'border-left:3px solid {accent};padding:10px 14px;'
                                  f'border-radius:6px;font-family:Courier New;font-size:12px;">'
                                  f'<code>')
                in_code = True
            continue
        if in_code:
            html_lines.append(line.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"))
            continue
        if line.startswith("#### "):
            html_lines.append(close_lists())
            html_lines.append(f'<h4 style="color:{accent};margin:10px 0 4px 0;">{_inline(line[5:])}</h4>')
            continue
        if line.startswith("### "):
            html_lines.append(close_lists())
            html_lines.append(f'<h3 style="color:{accent};margin:14px 0 6px 0;border-bottom:1px solid {accent}40;padding-bottom:4px;">{_inline(line[4:])}</h3>')
            continue
        if line.startswith("## "):
            html_lines.append(close_lists())
            html_lines.append(f'<h2 style="color:{accent};margin:18px 0 8px 0;border-bottom:2px solid {accent}60;padding-bottom:6px;font-size:15px;">{_inline(line[3:])}</h2>')
            continue
        if line.startswith("# "):
            html_lines.append(close_lists())
            html_lines.append(f'<h1 style="color:{accent};margin:20px 0 10px 0;font-size:18px;">{_inline(line[2:])}</h1>')
            continue
        if re.match(r"^[-*_]{3,}$", line.strip()):
            html_lines.append(close_lists())
            html_lines.append(f'<hr style="border:none;border-top:1px solid {accent}40;margin:12px 0;">')
            continue
        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                html_lines.append(close_lists())
                html_lines.append(f'<table style="border-collapse:collapse;width:100%;font-family:Courier New;font-size:12px;margin:10px 0;">')
                in_table = True
            if re.match(r"^\|[-| :]+\|$", line.strip()):
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if "td" not in "\n".join(html_lines[-2:]):
                row = "".join(f'<th style="padding:6px 10px;border:1px solid {accent}30;background:{accent}20;color:{accent};text-align:left;">{_inline(c)}</th>' for c in cells)
            else:
                row = "".join(f'<td style="padding:5px 10px;border:1px solid {accent}20;color:#b8c4d8;">{_inline(c)}</td>' for c in cells)
            html_lines.append(f"<tr>{row}</tr>")
            continue
        else:
            if in_table:
                html_lines.append("</table>"); in_table = False
        if line.startswith("> "):
            html_lines.append(close_lists())
            html_lines.append(f'<blockquote style="border-left:3px solid {accent};margin:6px 0;padding:4px 12px;color:#8899bb;background:{accent}10;border-radius:0 4px 4px 0;">{_inline(line[2:])}</blockquote>')
            continue
        if re.match(r"^[\s]*[-*+] ", line):
            if not in_ul:
                html_lines.append(close_lists())
                html_lines.append(f'<ul style="margin:4px 0;padding-left:20px;color:#b8c4d8;font-family:Courier New;font-size:13px;">')
                in_ul = True
            item = re.sub(r"^[\s]*[-*+] ", "", line)
            html_lines.append(f'<li style="margin:2px 0;">{_inline(item)}</li>')
            continue
        if re.match(r"^[\s]*\d+\. ", line):
            if not in_ol:
                html_lines.append(close_lists())
                html_lines.append(f'<ol style="margin:4px 0;padding-left:20px;color:#b8c4d8;font-family:Courier New;font-size:13px;">')
                in_ol = True
            item = re.sub(r"^[\s]*\d+\. ", "", line)
            html_lines.append(f'<li style="margin:2px 0;">{_inline(item)}</li>')
            continue
        else:
            html_lines.append(close_lists())
        if not line.strip():
            html_lines.append('<div style="margin:4px 0;"></div>')
            continue
        html_lines.append(f'<p style="margin:3px 0;color:#b8c4d8;font-family:Courier New;font-size:13px;line-height:1.6;">{_inline(line)}</p>')

    html_lines.append(close_lists())
    if in_table: html_lines.append("</table>")
    if in_code:  html_lines.append("</code></pre>")
    return "\n".join(html_lines)


def _inline(text: str) -> str:
    """Process inline markdown: bold, italic, code, etc."""
    text = text.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
    text = re.sub(r"\*\*\*(.+?)\*\*\*", r'<b><i>\1</i></b>', text)
    text = re.sub(r"\*\*(.+?)\*\*", r'<b style="color:#e2e8f0;">\1</b>', text)
    text = re.sub(r"__(.+?)__", r'<b style="color:#e2e8f0;">\1</b>', text)
    text = re.sub(r"\*(.+?)\*", r'<i style="color:#cbd5e1;">\1</i>', text)
    text = re.sub(r"`(.+?)`", r'<code style="background:#0d1117;color:#a8d8a8;padding:1px 5px;border-radius:3px;font-family:Courier New;font-size:12px;">\1</code>', text)
    text = re.sub(r"~~(.+?)~~", r'<s style="color:#555;">\1</s>', text)
    text = re.sub(r"==(.+?)==", r'<mark style="background:#f59e0b30;color:#fbbf24;padding:0 3px;">\1</mark>', text)
    return text

# ═══════════════════════════════════════════════════════════════════
#   PARTICLE SYSTEM
# ═══════════════════════════════════════════════════════════════════

class Particle:
    def __init__(self, w, h):
        self.reset(w, h)

    def reset(self, w, h):
        self.x  = random.uniform(0, w)
        self.y  = random.uniform(0, h)
        self.vx = random.uniform(-0.3, 0.3)
        self.vy = random.uniform(-0.5, -0.1)
        self.size   = random.uniform(1, 3)
        self.alpha  = random.uniform(0.1, 0.6)
        self.color  = random.choice(["#f59e0b","#60a5fa","#e879f9","#34d399","#a78bfa"])
        self.life   = random.uniform(0.5, 1.0)
        self.decay  = random.uniform(0.002, 0.006)

    def update(self, w, h):
        self.x += self.vx; self.y += self.vy
        self.life -= self.decay
        if self.life <= 0 or self.x < 0 or self.x > w or self.y < 0:
            self.reset(w, h)


class ParticleWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.particles = []
        self.setAttribute(Qt.WA_TransparentForMouseEvents)
        self.setAttribute(Qt.WA_NoSystemBackground)
        self.timer = QTimer(self)
        self.timer.timeout.connect(self._tick)
        self.timer.start(33)
        self.active = False

    def start(self):
        self.active = True
        if not self.particles:
            for _ in range(60):
                self.particles.append(Particle(max(self.width(),100), max(self.height(),100)))

    def stop(self):
        self.active = False

    def _tick(self):
        if self.active and self.particles:
            for p in self.particles:
                p.update(max(self.width(),100), max(self.height(),100))
            self.update()

    def paintEvent(self, _):
        if not self.active or not self.particles: return
        p = QPainter(self)
        p.setRenderHint(QPainter.Antialiasing)
        for pt in self.particles:
            col = QColor(pt.color)
            col.setAlphaF(pt.alpha * pt.life)
            p.setBrush(QBrush(col))
            p.setPen(Qt.NoPen)
            p.drawEllipse(QPointF(pt.x, pt.y), pt.size, pt.size)


# ═══════════════════════════════════════════════════════════════════
#   GLOWING PROGRESS BAR
# ═══════════════════════════════════════════════════════════════════

class GlowBar(QWidget):
    def __init__(self, color="#f59e0b", height=12, parent=None):
        super().__init__(parent)
        self.setFixedHeight(height)
        self._val   = 0
        self._color = QColor(color)
        self._pulse = 0.0
        self._dir   = 1
        self._t = QTimer(self)
        self._t.timeout.connect(self._tick)
        self._t.start(40)

    def set_value(self, v):
        self._val = max(0, min(100, v)); self.update()

    def _tick(self):
        self._pulse += 0.05 * self._dir
        if self._pulse >= 1.0: self._dir = -1
        if self._pulse <= 0.0: self._dir = 1
        self.update()

    def paintEvent(self, _):
        p = QPainter(self)
        p.setRenderHint(QPainter.Antialiasing)
        w, h = self.width(), self.height()
        r = h // 2
        p.setBrush(QBrush(QColor("#111128")))
        p.setPen(Qt.NoPen)
        p.drawRoundedRect(0, 0, w, h, r, r)
        fw = int(w * self._val / 100)
        if fw > 0:
            grad = QLinearGradient(0, 0, fw, 0)
            c1 = self._color; c2 = QColor(self._color)
            grad.setColorAt(0, c1); grad.setColorAt(1, c2)
            p.setBrush(QBrush(grad))
            p.drawRoundedRect(0, 0, fw, h, r, r)
            alpha = int(60 + 40 * self._pulse)
            glow_c = QColor(self._color); glow_c.setAlpha(alpha)
            gp = QPen(glow_c, 4)
            p.setPen(gp); p.setBrush(Qt.NoBrush)
            p.drawRoundedRect(1, 1, fw-2, h-2, r, r)


# ═══════════════════════════════════════════════════════════════════
#   SCORE RING
# ═══════════════════════════════════════════════════════════════════

class ScoreRing(QWidget):
    def __init__(self, label="SCORE", color="#e879f9", size=100):
        super().__init__()
        self.setFixedSize(size, size)
        self._val   = 0.0
        self._label = label
        self._color = QColor(color)
        self._anim  = 0.0
        self._t = QTimer(self); self._t.timeout.connect(self._tick); self._t.start(40)

    def set_value(self, v):
        self._val = max(0, min(10, v)); self.update()

    def _tick(self):
        self._anim += 0.03
        self.update()

    def paintEvent(self, _):
        p = QPainter(self)
        p.setRenderHint(QPainter.Antialiasing)
        w, h = self.width(), self.height()
        cx, cy = w//2, h//2
        r = min(w,h)//2 - 8
        p.setPen(QPen(QColor("#1a1a30"), 6))
        p.setBrush(Qt.NoBrush)
        p.drawEllipse(QPointF(cx, cy), r, r)
        if self._val > 0:
            span = int(self._val / 10 * 360 * 16)
            glow_alpha = int(180 + 60 * math.sin(self._anim))
            gc = QColor(self._color); gc.setAlpha(glow_alpha)
            p.setPen(QPen(gc, 10, Qt.SolidLine, Qt.RoundCap))
            p.drawArc(QRectF(cx-r, cy-r, r*2, r*2), 90*16, -span)
            p.setPen(QPen(self._color, 6, Qt.SolidLine, Qt.RoundCap))
            p.drawArc(QRectF(cx-r, cy-r, r*2, r*2), 90*16, -span)
        p.setPen(self._color)
        p.setFont(QFont("Courier New", 18, QFont.Bold))
        score_txt = f"{self._val:.1f}" if self._val > 0 else "—"
        p.drawText(QRectF(0,0,w,h-14), Qt.AlignCenter, score_txt)
        p.setFont(QFont("Courier New", 7))
        p.setPen(QColor("#444"))
        p.drawText(QRectF(0, cy+r//2, w, 18), Qt.AlignCenter, self._label)


# ═══════════════════════════════════════════════════════════════════
#   SCORE GRAPH
# ═══════════════════════════════════════════════════════════════════

class ScoreGraph(QWidget):
    def __init__(self):
        super().__init__()
        self.scores = []; self.setMinimumHeight(150)
        self._anim = 0.0
        t = QTimer(self); t.timeout.connect(self._tick); t.start(50)

    def _tick(self): self._anim += 0.04; self.update()

    def add_score(self, s): self.scores.append(s); self.update()

    def clear(self): self.scores = []; self.update()

    def paintEvent(self, _):
        p = QPainter(self); p.setRenderHint(QPainter.Antialiasing)
        w, h, pad = self.width(), self.height(), 34
        grad = QLinearGradient(0,0,0,h)
        grad.setColorAt(0, QColor("#0a0a18")); grad.setColorAt(1, QColor("#070710"))
        p.fillRect(0, 0, w, h, QBrush(grad))
        for i in range(0, 11, 2):
            y = pad + (h-2*pad)*(1-i/10)
            p.setPen(QPen(QColor("#151528"), 1, Qt.DotLine))
            p.drawLine(pad, int(y), w-8, int(y))
            p.setPen(QColor("#333")); p.setFont(QFont("Courier New",7))
            p.drawText(2, int(y)+4, str(i))
        yt = pad+(h-2*pad)*(1-8/10)
        glow_a = int(40 + 20*math.sin(self._anim))
        p.setPen(QPen(QColor(34,197,94, glow_a), 1, Qt.DashLine))
        p.drawLine(pad, int(yt), w-8, int(yt))
        if not self.scores:
            p.setPen(QColor("#252540")); p.setFont(QFont("Courier New",10))
            p.drawText(w//2-55, h//2+4, "Awaiting results..."); return
        n    = len(self.scores)
        step = (w-pad-12)/max(n-1,1) if n>1 else 0
        pts  = [(pad+i*step, pad+(h-2*pad)*(1-s/10)) for i,s in enumerate(self.scores)]
        if len(pts) > 1:
            path = QPainterPath()
            path.moveTo(pts[0][0], h-pad)
            for x,y in pts: path.lineTo(x, y)
            path.lineTo(pts[-1][0], h-pad); path.closeSubpath()
            ag = QLinearGradient(0,pad,0,h-pad)
            ag.setColorAt(0, QColor(232,121,249,70)); ag.setColorAt(1, QColor(232,121,249,0))
            p.fillPath(path, QBrush(ag))
        p.setPen(QPen(QColor("#e879f9"), 2, Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin))
        for i in range(1, len(pts)):
            p.drawLine(QPointF(*pts[i-1]), QPointF(*pts[i]))
        for i,(x,y) in enumerate(pts):
            s = self.scores[i]
            c = QColor("#22c55e" if s>=8 else "#f59e0b" if s>=6.5 else "#f87171")
            is_last = (i == len(pts)-1)
            pulse_r = 5 + (2*math.sin(self._anim*2)) if is_last else 5
            gc = QColor(c); gc.setAlpha(40)
            p.setBrush(QBrush(gc)); p.setPen(Qt.NoPen)
            p.drawEllipse(QPointF(x,y), pulse_r+4, pulse_r+4)
            p.setBrush(QBrush(c)); p.setPen(QPen(QColor("#000"),1))
            p.drawEllipse(QPointF(x,y), pulse_r, pulse_r)
            p.setPen(c); p.setFont(QFont("Courier New",8,QFont.Bold))
            p.drawText(int(x)-12, int(y)-10, f"{s:.1f}")

# ═══════════════════════════════════════════════════════════════════
#   KNOWLEDGE GRAPH WIDGET
# ═══════════════════════════════════════════════════════════════════

class KnowledgeGraph(QWidget):
    def __init__(self):
        super().__init__()
        self.setMinimumHeight(200)
        self.nodes = []
        self.edges = []
        self._anim = 0.0
        t = QTimer(self); t.timeout.connect(self._tick); t.start(40)

    def _tick(self): self._anim += 0.02; self.update()

    def set_topic(self, topic):
        words = [w for w in topic.split() if len(w) > 3][:6]
        w, h  = max(self.width(),200), max(self.height(),200)
        self.nodes = [{"id":0,"label":topic[:18]+"…" if len(topic)>18 else topic,
                       "x":w//2,"y":h//2,"color":"#f59e0b","size":14}]
        self.edges = []
        colors = ["#60a5fa","#e879f9","#34d399","#fb923c","#a78bfa","#38bdf8"]
        angle_step = 2*math.pi/max(len(words),1)
        for i,w_text in enumerate(words):
            angle = i*angle_step
            rx = w//2 + int(math.cos(angle)*(w//3))
            ry = h//2 + int(math.sin(angle)*(h//3))
            nid = i+1
            self.nodes.append({"id":nid,"label":w_text,"x":rx,"y":ry,
                               "color":colors[i%len(colors)],"size":9})
            self.edges.append({"src":0,"dst":nid})
        self.update()

    def paintEvent(self, _):
        p = QPainter(self); p.setRenderHint(QPainter.Antialiasing)
        w, h = self.width(), self.height()
        p.fillRect(0,0,w,h, QColor("#090915"))
        if not self.nodes: return
        for e in self.edges:
            src = next((n for n in self.nodes if n["id"]==e["src"]),None)
            dst = next((n for n in self.nodes if n["id"]==e["dst"]),None)
            if src and dst:
                glow_a = int(50 + 30*math.sin(self._anim + dst["id"]))
                p.setPen(QPen(QColor(96,165,250,glow_a), 1))
                p.drawLine(QPointF(src["x"],src["y"]), QPointF(dst["x"],dst["y"]))
        for n in self.nodes:
            pulse = math.sin(self._anim*1.5 + n["id"]*0.8)
            extra = 2*pulse if n["id"]==0 else 0
            r = n["size"] + extra
            c = QColor(n["color"])
            gc = QColor(c); gc.setAlpha(40)
            p.setBrush(QBrush(gc)); p.setPen(Qt.NoPen)
            p.drawEllipse(QPointF(n["x"],n["y"]), r+6, r+6)
            p.setBrush(QBrush(c)); p.setPen(QPen(QColor("#000"),1))
            p.drawEllipse(QPointF(n["x"],n["y"]), r, r)
            p.setPen(c); p.setFont(QFont("Courier New",8))
            p.drawText(int(n["x"])-40, int(n["y"])+int(r)+12, 80, 14,
                       Qt.AlignCenter, n["label"])


# ═══════════════════════════════════════════════════════════════════
#   ACHIEVEMENT BANNER
# ═══════════════════════════════════════════════════════════════════

class AchievementBanner(QFrame):
    def __init__(self):
        super().__init__()
        self.setFixedHeight(0)
        self.setStyleSheet("background:#111; border-radius:6px; border:1px solid #333;")
        self._lay = QHBoxLayout(self)
        self._lay.setContentsMargins(12,4,12,4)
        self._icon = QLabel("🏆")
        self._icon.setFont(QFont("Segoe UI Emoji",14))
        self._icon.setStyleSheet("background:transparent;border:none;")
        self._text = QLabel("")
        self._text.setFont(QFont("Courier New",11,QFont.Bold))
        self._text.setStyleSheet("background:transparent;border:none;")
        self._lay.addWidget(self._icon)
        self._lay.addWidget(self._text)
        self._lay.addStretch()
        self._anim = QPropertyAnimation(self, b"minimumHeight")
        self._anim.setDuration(300)
        self._anim.setEasingCurve(QEasingCurve.OutBack)
        self._hide_timer = QTimer(); self._hide_timer.setSingleShot(True)
        self._hide_timer.timeout.connect(self.hide_banner)

    def show_achievement(self, icon, text, color="#f59e0b"):
        self._icon.setText(icon)
        self._text.setText(text)
        self._text.setStyleSheet(f"color:{color};background:transparent;border:none;")
        self.setStyleSheet(f"background:#0d0d18;border-radius:6px;border:1px solid {color}60;")
        self._anim.setStartValue(0); self._anim.setEndValue(38)
        self._anim.start()
        self._hide_timer.start(3500)

    def hide_banner(self):
        self._anim.setStartValue(38); self._anim.setEndValue(0)
        self._anim.start()


# ═══════════════════════════════════════════════════════════════════
#   CHAT BUBBLE (with markdown rendering)
# ═══════════════════════════════════════════════════════════════════

AGENT_ACCENTS = {
    "director":"#f59e0b","literature":"#38bdf8","hypothesis":"#fbbf24",
    "writer":"#60a5fa","reviewer_methods":"#f87171","reviewer_novelty":"#c084fc",
    "reviewer_harsh":"#fb923c","defender":"#34d399","corrector":"#4ade80",
    "humanizer":"#a78bfa","evaluator":"#e879f9",
}

class ChatBubble(QFrame):
    def __init__(self, agent_id, iteration=None):
        super().__init__()
        ag  = AGENTS_MAP[agent_id]
        clr = ag["color"]
        iter_tag = f"  ·  Iter {iteration}" if iteration else ""
        self.setObjectName("bub")
        self.setStyleSheet(f"""
          QFrame#bub{{background:qlineargradient(x1:0,y1:0,x2:0,y2:1,
              stop:0 #0e0e20,stop:1 #090914);
            border:1px solid {clr}40;border-left:3px solid {clr};
            border-radius:10px;margin:3px 8px;}}
        """)
        lay = QVBoxLayout(self); lay.setContentsMargins(14,10,14,10); lay.setSpacing(6)
        hdr = QHBoxLayout()
        icon = QLabel(ag["icon"]); icon.setFont(QFont("Segoe UI Emoji",15))
        icon.setStyleSheet("background:transparent;border:none;")
        icon.setFixedWidth(26)
        name = QLabel(ag["name"]+iter_tag)
        name.setFont(QFont("Courier New",11,QFont.Bold))
        name.setStyleSheet(f"color:{clr};background:transparent;border:none;")
        role = QLabel(f"  /  {ag['role']}")
        role.setFont(QFont("Courier New",9))
        role.setStyleSheet("color:#2a2a4a;background:transparent;border:none;")
        self.ts = QLabel("")
        self.ts.setFont(QFont("Courier New",8))
        self.ts.setStyleSheet("color:#1e1e3a;background:transparent;border:none;")
        self.ts.setAlignment(Qt.AlignRight)
        hdr.addWidget(icon); hdr.addWidget(name); hdr.addWidget(role)
        hdr.addStretch(); hdr.addWidget(self.ts)
        lay.addLayout(hdr)
        self.tabs = QTabWidget()
        self.tabs.setStyleSheet(f"""
          QTabWidget::pane{{border:1px solid {clr}20;border-radius:4px;background:#08080f;}}
          QTabBar::tab{{background:#0c0c1e;color:#444;font-family:Courier New;font-size:9px;
            padding:3px 10px;border-radius:3px 3px 0 0;}}
          QTabBar::tab:selected{{background:#111128;color:{clr};}}
        """)
        self.tabs.setMinimumHeight(60)
        self.tabs.setMaximumHeight(420)
        self.md_view = QTextBrowser()
        self.md_view.setFont(QFont("Courier New",11))
        self.md_view.setStyleSheet(f"""
          QTextBrowser{{background:#08080f;color:#b8c4d8;border:none;
            selection-background-color:{clr}30;}}
          QScrollBar:vertical{{background:#0a0a18;width:5px;border-radius:2px;}}
          QScrollBar::handle:vertical{{background:#1e1e3a;border-radius:2px;}}
        """)
        self.md_view.setOpenExternalLinks(True)
        self.raw_view = QPlainTextEdit()
        self.raw_view.setReadOnly(True)
        self.raw_view.setFont(QFont("Courier New",10))
        self.raw_view.setStyleSheet(f"""
          QPlainTextEdit{{background:#06060e;color:#6a7a9a;border:none;}}
          QScrollBar:vertical{{background:#0a0a18;width:5px;border-radius:2px;}}
          QScrollBar::handle:vertical{{background:#1e1e3a;border-radius:2px;}}
        """)
        self.tabs.addTab(self.md_view, "◈ RENDERED")
        self.tabs.addTab(self.raw_view, "⟨⟩ RAW")
        lay.addWidget(self.tabs)
        self.status_bar = QHBoxLayout()
        self.status_lbl = QLabel("● STREAMING...")
        self.status_lbl.setFont(QFont("Courier New",9))
        self.status_lbl.setStyleSheet(f"color:{clr};background:transparent;border:none;")
        self.wc_lbl = QLabel("")
        self.wc_lbl.setFont(QFont("Courier New",8))
        self.wc_lbl.setStyleSheet("color:#2a2a4a;background:transparent;border:none;")
        self.wc_lbl.setAlignment(Qt.AlignRight)
        self.status_bar.addWidget(self.status_lbl)
        self.status_bar.addStretch()
        self.status_bar.addWidget(self.wc_lbl)
        lay.addLayout(self.status_bar)
        self._raw   = ""
        self._color = clr
        self._agent = ag

    def append(self, tok):
        self._raw += tok
        wc = len(self._raw.split())
        self.wc_lbl.setText(f"{wc} words")
        self.raw_view.setPlainText(self._raw)
        cr = self.raw_view.textCursor(); cr.movePosition(QTextCursor.End)
        self.raw_view.setTextCursor(cr)
        if len(self._raw) % 60 == 0 or len(tok) > 20:
            self._render_md()
            sb = self.md_view.verticalScrollBar()
            sb.setValue(sb.maximum())

    def _render_md(self):
        html = md_to_html(self._raw, self._color)
        self.md_view.setHtml(f"""
          <html><body style="background:#08080f;padding:8px;margin:0;">
          {html}</body></html>""")

    def done(self):
        self._render_md()
        sb = self.md_view.verticalScrollBar(); sb.setValue(sb.maximum())
        wc = len(self._raw.split())
        self.ts.setText(datetime.now().strftime("%H:%M:%S"))
        self.status_lbl.setText("✓ COMPLETE")
        self.status_lbl.setStyleSheet(
            "color:#22c55e;background:transparent;border:none;font-weight:bold;")
        self.wc_lbl.setText(f"{wc} words")
        h = min(int(self.md_view.document().size().height())+30, 420)
        self.tabs.setMaximumHeight(h)

    def error(self, msg):
        self.status_lbl.setText(f"✗ {msg[:100]}")
        self.status_lbl.setStyleSheet("color:#f87171;background:transparent;border:none;")

# ═══════════════════════════════════════════════════════════════════
#   AGENT + PIPELINE DEFINITIONS
# ═══════════════════════════════════════════════════════════════════

AGENTS = [
  {"id":"director","name":"Research Director","role":"Topic Strategist","icon":"◈","color":"#f59e0b",
   "system":"""You are a Research Director AI. Output detailed markdown with:
## RESEARCH DIRECTION
Core Problem: [one crisp sentence]
Paper Type: [as specified]
## Paper Outline
1. Abstract  2. Introduction  3. Background & Related Work
4. Methodology  5. Experiments & Results  6. Discussion
7. Conclusion  8. References
## Research Sub-Questions
- [4-6 focused sub-questions]
## Keywords
[6-10 comma-separated keywords]
## Expected Contribution
[3-4 sentences on novelty]
## Recommended Baselines
- [3-5 systems to compare against]"""},

  {"id":"literature","name":"Literature Intelligence","role":"Scholar Agent","icon":"📚","color":"#38bdf8",
   "system":"""You are a Literature Intelligence Agent. Output detailed markdown.
## Literature Analysis
### Key Related Works
| Paper | Year | Core Contribution | Limitation |
|-------|------|-------------------|------------|
(List 8-12 papers with realistic titles)
### Research Gap Identified
[2-3 paragraphs on what is MISSING]
### Citation List
[1] Author et al., "Title," Venue, Year.
(10-15 references in specified citation style)"""},

  {"id":"hypothesis","name":"Hypothesis Generator","role":"Scientific Reasoner","icon":"💡","color":"#fbbf24",
   "system":"""You are a Scientific Hypothesis Generator. Output detailed markdown.
## Scientific Hypotheses
H1: [Primary hypothesis]  H2: [Secondary hypothesis]  H3: [Alternative]
## Experiment Design
### Datasets
- Dataset 1: [name, description, size, why chosen]
### Baselines
- Baseline 1: [why relevant]
### Evaluation Metrics
| Metric | Description | Target |
|--------|-------------|--------|
### Expected Results
| Method | Metric 1 | Metric 2 |
|--------|----------|----------|"""},

  {"id":"writer","name":"Academic Writer","role":"Paper Generator","icon":"✍","color":"#60a5fa",
   "system":"""You are an Academic Writer AI. Write a complete full-length research paper in markdown.
Include ALL sections: Abstract (150-200 words), Introduction, Background & Related Work,
Methodology (with equations in $...$ notation), Experiments & Results (with tables),
Discussion, Conclusion, References.
Rules: minimum 2500 words, technical and precise, address ALL previous reviewer feedback."""},

  {"id":"reviewer_methods","name":"Methods Reviewer","role":"Methodology Expert","icon":"⚗","color":"#f87171",
   "system":"""You are a Methodology Expert Reviewer. Output structured markdown review.
## Methodology Review
### Algorithm Correctness
### Equation Consistency
### Experimental Validity
| Criterion | Assessment | Score |
|-----------|------------|-------|
### Missing Experiments
### Technical Flaws
Methods Score: X/10"""},

  {"id":"reviewer_novelty","name":"Novelty Reviewer","role":"Originality Expert","icon":"🔭","color":"#c084fc",
   "system":"""You are a Novelty and Originality Reviewer.
## Novelty Review
### Core Contribution Assessment
### Overlap Analysis
| Related Work | Overlap Level | Key Difference |
### Novelty Scores
| Dimension | Score | Comment |
Novelty Score: X/10"""},

  {"id":"reviewer_harsh","name":"Harsh Reviewer","role":"Conference Gatekeeper","icon":"⚡","color":"#fb923c",
   "system":"""You are a HARSH NeurIPS/ICML/ICLR reviewer. Be brutally specific.
## Harsh Conference Review
Decision: [Strong Accept / Accept / Weak Accept / Weak Reject / Reject]
### Strengths
### Critical Weaknesses
> ⚠️ Major: [devastating critique]
### Questions for Authors
Overall Score: X/10
Confidence: [Expert / High / Medium / Low]"""},

  {"id":"defender","name":"Argument Defender","role":"Author Rebuttal","icon":"🛡","color":"#34d399",
   "system":"""You are an Author Rebuttal Agent.
## Author Rebuttal
### Response to Methods Reviewer
### Response to Novelty Reviewer
### Response to Harsh Reviewer
### Points Conceded ✓ (will fix in revision)
### Points Defended ✗ (maintained)
### Summary"""},

  {"id":"corrector","name":"Correction Engine","role":"Technical Editor","icon":"⚙","color":"#4ade80",
   "system":"""You are a Technical Correction Engine.
Fix EVERY conceded point from the rebuttal and all reviewer concerns.
Write the COMPLETE corrected paper in markdown.
Mark changes: [FIXED: reason] inline.
Target 2500+ words."""},

  {"id":"humanizer","name":"Humanization Engine","role":"Writing Naturalizer","icon":"🧬","color":"#a78bfa",
   "system":"""You are a Humanization Engine. Transform AI-generated text to natural researcher writing.
Rules:
1. Replace AI patterns (e.g. "This paper presents" → "In this work, we investigate")
2. Add researcher voice: "we believe", "our intuition suggests", "somewhat surprisingly"
3. Vary sentence structure
4. Preserve ALL technical content exactly
Output the COMPLETE humanized paper in clean markdown."""},

  {"id":"evaluator","name":"Journal Evaluator","role":"Quality Gate","icon":"★","color":"#e879f9",
   "system":"""You are a Journal Quality Evaluator.
## Final Evaluation
### Quality Scorecard
| Criterion | Score | Comment |
|-----------|-------|---------|
| Novelty | X/10 | |
| Technical Rigor | X/10 | |
| Experimental Validation | X/10 | |
| Writing Quality | X/10 | |
| Literature Coverage | X/10 | |
| Reproducibility | X/10 | |
### Publication Score: X.X/10
### VERDICT
> ✅ APPROVED / ⚠️ CONDITIONAL / ❌ NEEDS_REVISION
> Score MUST be X.X format like 7.8/10"""},
]

AGENTS_MAP = {a["id"]: a for a in AGENTS}
REVIEWER_IDS = ["reviewer_methods","reviewer_novelty","reviewer_harsh"]
PAPER_TYPES    = ["Original Research Paper","Survey Paper","Literature Review",
                  "Short Paper / Extended Abstract","Thesis Chapter","Research Proposal"]
CITATION_STYLES = ["IEEE","APA","ACM","Springer","NeurIPS"]
MODELS = ["gpt-4o","gpt-4o-mini","gpt-4-turbo","gpt-3.5-turbo"]
ACHIEVEMENTS = [
    (7.0, "🌟", "GOOD PAPER", "#f59e0b"),
    (8.0, "🏆", "PUBLISHABLE QUALITY!", "#22c55e"),
    (8.5, "🚀", "CONFERENCE READY!", "#60a5fa"),
    (9.0, "💎", "OUTSTANDING RESEARCH!", "#e879f9"),
    (9.5, "👑", "WORLD CLASS PAPER!", "#fbbf24"),
]

# ═══════════════════════════════════════════════════════════════════
#   STREAMING WORKER
# ═══════════════════════════════════════════════════════════════════

class AgentWorker(QThread):
    token_received = Signal(str, str)
    agent_done     = Signal(str, str)
    agent_error    = Signal(str, str)
    score_found    = Signal(float)

    def __init__(self, agent_id, messages, api_key, model):
        super().__init__()
        self.agent_id = agent_id; self.messages = messages
        self.api_key = api_key;   self.model = model
        self._stop = False

    def stop(self): self._stop = True

    def run(self):
        try:
            hdrs = {"Content-Type":"application/json",
                    "Authorization":f"Bearer {self.api_key}"}
            body = {"model":self.model,"messages":self.messages,
                    "max_tokens":3500,"temperature":0.75,"stream":True}
            full = ""
            with requests.post("https://api.openai.com/v1/chat/completions",
                               headers=hdrs,json=body,stream=True,timeout=180) as r:
                if r.status_code != 200:
                    e = r.json(); self.agent_error.emit(self.agent_id,
                        e.get("error",{}).get("message",str(e))); return
                for line in r.iter_lines():
                    if self._stop: break
                    if not line: continue
                    line = line.decode("utf-8")
                    if line.startswith("data: "):
                        d = line[6:]
                        if d=="[DONE]": break
                        try:
                            tok = json.loads(d)["choices"][0]["delta"].get("content","")
                            if tok: full+=tok; self.token_received.emit(self.agent_id,tok)
                        except: pass
            if self.agent_id=="evaluator":
                m = re.search(r"(\d+\.?\d*)\s*/\s*10", full)
                if m: self.score_found.emit(float(m.group(1)))
            self.agent_done.emit(self.agent_id, full)
        except Exception as e:
            self.agent_error.emit(self.agent_id, str(e))


# ═══════════════════════════════════════════════════════════════════
#   SECTION DIVIDER
# ═══════════════════════════════════════════════════════════════════

class IterDiv(QFrame):
    def __init__(self, label):
        super().__init__()
        self.setStyleSheet("background:#0e0e28;border-radius:6px;margin:12px 6px 4px 6px;"
                           "border:1px solid #1e1e48;")
        l = QHBoxLayout(self); l.setContentsMargins(16,8,16,8)
        lbl = QLabel(label); lbl.setFont(QFont("Courier New",10,QFont.Bold))
        lbl.setStyleSheet("color:#3a3a70;background:transparent;border:none;letter-spacing:3px;")
        l.addWidget(lbl); l.addStretch()

# ═══════════════════════════════════════════════════════════════════
#   DOCX EXPORTER
# ═══════════════════════════════════════════════════════════════════

def export_docx(filepath, topic, paper_md, director, literature, hypothesis,
                reviews, final_eval, best_iter, best_score, timestamp):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    def add_heading(text, level=1, color=(31,78,121)):
        p = doc.add_heading(text, level=level)
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = RGBColor(*color)
        run.bold = True
        return p

    def md_to_docx(text):
        lines = text.split("\n")
        for raw in lines:
            line = raw
            if not line.strip(): doc.add_paragraph(); continue
            if line.startswith("# "):
                p = doc.add_heading(line[2:], level=1)
                if p.runs: p.runs[0].font.color.rgb = RGBColor(31,78,121)
            elif line.startswith("## "):
                p = doc.add_heading(line[3:], level=2)
                if p.runs: p.runs[0].font.color.rgb = RGBColor(31,100,140)
            elif line.startswith("### "):
                p = doc.add_heading(line[4:], level=3)
                if p.runs: p.runs[0].font.color.rgb = RGBColor(60,120,160)
            elif line.startswith("> "):
                p = doc.add_paragraph(style="Quote")
                p.add_run(line[2:]).italic = True
            elif re.match(r"^[-*] ", line):
                doc.add_paragraph(re.sub(r"^[-*] ","",line).strip(), style="List Bullet")
            elif re.match(r"^\d+\. ", line):
                doc.add_paragraph(re.sub(r"^\d+\. ","",line).strip(), style="List Number")
            elif "|" in line and line.strip().startswith("|"):
                pass  # Skip tables in simple mode
            else:
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
                clean = re.sub(r"\*(.+?)\*", r"\1", clean)
                clean = re.sub(r"`(.+?)`", r"\1", clean)
                p = doc.add_paragraph(clean)
                p.paragraph_format.space_after = Pt(3)

    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(topic)
    title_run.font.size = Pt(22); title_run.bold = True
    title_run.font.color.rgb = RGBColor(30,100,180)
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(f"Generated by Life Of Research  ·  By Pranay Mahendrakar  ·  {timestamp}\n"
                      f"Best Score: {best_score:.1f}/10  ·  Iteration {best_iter}")
    mr.font.size = Pt(10); mr.italic = True
    mr.font.color.rgb = RGBColor(120,120,140)
    doc.add_page_break()

    add_heading("Research Direction & Plan", 1, (31,78,121))
    md_to_docx(director)
    doc.add_page_break()

    add_heading("Literature Analysis", 1, (31,78,121))
    md_to_docx(literature)
    doc.add_page_break()

    add_heading("Hypotheses & Experiment Design", 1, (31,78,121))
    md_to_docx(hypothesis)
    doc.add_page_break()

    add_heading("Research Paper", 1, (10,80,140))
    md_to_docx(paper_md)
    doc.add_page_break()

    add_heading("Peer Review Panel", 1, (100,30,30))
    for rid, rtxt in reviews.items():
        ag = AGENTS_MAP.get(rid, {})
        add_heading(f"{ag.get('name','Review')} — {ag.get('role','')}", 2, (140,50,50))
        md_to_docx(rtxt)
    doc.add_page_break()

    add_heading("Final Journal Evaluation", 1, (80,30,120))
    md_to_docx(final_eval)

    doc.save(filepath)

# ═══════════════════════════════════════════════════════════════════
#   MAIN WINDOW
# ═══════════════════════════════════════════════════════════════════

class LifeOfResearchWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Life Of Research  ·  By Pranay Mahendrakar")
        self.setMinimumSize(1400, 880)
        self.resize(1600, 960)
        # State
        self.worker       = None
        self.running      = False
        self.abort_flag   = False
        self.iteration    = 0
        self.scores       = []
        self.memory       = {}
        self.outputs      = {}
        self.bubbles      = {}
        self.best_paper   = ""
        self.best_score   = 0.0
        self.best_iter    = 0
        self._ach_shown   = set()
        self._total_words = 0
        self._ticker_msgs = [
            "◈ Life Of Research  ·  By Pranay Mahendrakar  ·  14-Agent Pipeline",
            "💡 Hypothesis Generator  ·  Literature Intelligence  ·  Debate Mode",
            "📚 IEEE  ·  APA  ·  ACM  ·  Springer  ·  NeurIPS citation styles",
            "★ Journal-quality evaluation with iterative improvement loop",
            "🧬 Humanization Engine removes AI patterns from writing",
            "⬇ Export to DOCX  ·  Markdown  ·  HTML with full formatting",
        ]
        self._ticker_idx = 0
        self._setup_palette()
        self._build_ui()

    def _setup_palette(self):
        pal = QPalette()
        pal.setColor(QPalette.Window,         QColor("#07070d"))
        pal.setColor(QPalette.WindowText,      QColor("#c0cce0"))
        pal.setColor(QPalette.Base,            QColor("#09091a"))
        pal.setColor(QPalette.Text,            QColor("#c0cce0"))
        pal.setColor(QPalette.Button,          QColor("#0e0e1e"))
        pal.setColor(QPalette.ButtonText,      QColor("#c0cce0"))
        pal.setColor(QPalette.Highlight,       QColor("#f59e0b"))
        pal.setColor(QPalette.HighlightedText, QColor("#000"))
        QApplication.setPalette(pal)

    def _build_ui(self):
        root = QWidget(); self.setCentralWidget(root)
        vbox = QVBoxLayout(root)
        vbox.setContentsMargins(0,0,0,0); vbox.setSpacing(0)
        vbox.addWidget(self._make_header())
        self.achievement = AchievementBanner()
        vbox.addWidget(self.achievement)
        split = QSplitter(Qt.Horizontal)
        split.setHandleWidth(3)
        split.setStyleSheet("QSplitter::handle{background:#101028;}")
        split.addWidget(self._make_left())
        split.addWidget(self._make_center())
        split.addWidget(self._make_right())
        split.setSizes([300,840,330])
        vbox.addWidget(split, 1)
        vbox.addWidget(self._make_statusbar())

    def _make_header(self):
        h = QFrame(); h.setFixedHeight(64)
        h.setStyleSheet("background:#05050f;border-bottom:1px solid #10103a;")
        lay = QHBoxLayout(h); lay.setContentsMargins(20,0,20,0)
        logo = QLabel("◈"); logo.setFont(QFont("Courier New",26,QFont.Bold))
        logo.setStyleSheet("color:#f59e0b;")
        title = QLabel(); title.setTextFormat(Qt.RichText)
        title.setText("<span style='color:#fff;font-family:Courier New;font-size:18px;font-weight:900;'>"
                      "LIFE OF <span style='color:#f59e0b'>RESEARCH</span></span>")
        badge = QLabel("v5.0")
        badge.setStyleSheet("background:#f59e0b;color:#000;padding:2px 8px;"
                            "border-radius:10px;font-family:Courier New;font-size:9px;font-weight:bold;")
        self.ticker_lbl = QLabel(self._ticker_msgs[0])
        self.ticker_lbl.setFont(QFont("Courier New",9))
        self.ticker_lbl.setStyleSheet("color:#2a2a5a;")
        self._ticker_t = QTimer(); self._ticker_t.timeout.connect(self._advance_ticker)
        self._ticker_t.start(4000)
        self.timer_lbl = QLabel("00:00:00")
        self.timer_lbl.setFont(QFont("Courier New",12,QFont.Bold))
        self.timer_lbl.setStyleSheet("color:#1e1e48;")
        self._elapsed = 0; self._session_t = QTimer()
        self._session_t.timeout.connect(self._tick_timer)
        self._session_t.start(1000)
        lay.addWidget(logo); lay.addSpacing(8); lay.addWidget(title)
        lay.addSpacing(8);  lay.addWidget(badge)
        lay.addSpacing(24); lay.addWidget(self.ticker_lbl)
        lay.addStretch();   lay.addWidget(self.timer_lbl)
        return h

    def _advance_ticker(self):
        self._ticker_idx = (self._ticker_idx+1) % len(self._ticker_msgs)
        self.ticker_lbl.setText(self._ticker_msgs[self._ticker_idx])

    def _tick_timer(self):
        if self.running: self._elapsed += 1
        h,m,s = self._elapsed//3600, (self._elapsed%3600)//60, self._elapsed%60
        self.timer_lbl.setText(f"{h:02d}:{m:02d}:{s:02d}")
        color = "#f59e0b" if self.running else "#1e1e48"
        self.timer_lbl.setStyleSheet(f"color:{color};font-family:Courier New;font-size:12px;font-weight:bold;")

    def _make_left(self):
        w = QWidget(); w.setMaximumWidth(320)
        w.setStyleSheet("background:#07070e;")
        lay = QVBoxLayout(w); lay.setContentsMargins(14,14,14,14); lay.setSpacing(10)

        def sec(txt, c="#f59e0b"):
            l = QLabel(txt); l.setFont(QFont("Courier New",8,QFont.Bold))
            l.setStyleSheet(f"color:{c};letter-spacing:2px;background:transparent;")
            return l

        def inp(ph="",pw=False):
            e = QLineEdit(); e.setPlaceholderText(ph)
            if pw: e.setEchoMode(QLineEdit.Password)
            e.setFont(QFont("Courier New",11))
            e.setStyleSheet("""QLineEdit{background:#0c0c1e;color:#c0cce0;
              border:1px solid #1a1a38;border-radius:5px;padding:8px 10px;}
              QLineEdit:focus{border-color:#f59e0b;}""")
            return e

        def cmb(items):
            c = QComboBox(); c.addItems(items); c.setFont(QFont("Courier New",10))
            c.setStyleSheet("""QComboBox{background:#0c0c1e;color:#c0cce0;
              border:1px solid #1a1a38;border-radius:5px;padding:6px 10px;}
              QComboBox::drop-down{border:none;}
              QComboBox QAbstractItemView{background:#0c0c1e;color:#c0cce0;
                border:1px solid #1a1a38;selection-background-color:#f59e0b30;}""")
            return c

        lay.addWidget(sec("◈ API KEY"))
        self.api_inp = inp("sk-...", pw=True); lay.addWidget(self.api_inp)
        lay.addWidget(sec("◈ MODEL"))
        self.model_cb = cmb(MODELS); lay.addWidget(self.model_cb)
        lay.addWidget(sec("◈ RESEARCH TOPIC"))
        self.topic_edit = QTextEdit()
        self.topic_edit.setPlaceholderText("e.g. Efficient KV Cache Optimization in Large Language Models")
        self.topic_edit.setFont(QFont("Courier New",10))
        self.topic_edit.setMaximumHeight(80)
        self.topic_edit.setStyleSheet("""QTextEdit{background:#0c0c1e;color:#c0cce0;
          border:1px solid #1a1a38;border-radius:5px;padding:8px;}
          QTextEdit:focus{border-color:#f59e0b;}""")
        lay.addWidget(self.topic_edit)
        row = QHBoxLayout()
        self.type_cb = cmb(PAPER_TYPES)
        self.cite_cb = cmb(CITATION_STYLES)
        row.addWidget(self.type_cb,3); row.addWidget(self.cite_cb,2)
        lay.addWidget(sec("◈ PAPER TYPE  /  CITATION")); lay.addLayout(row)
        lay.addWidget(sec("◈ LOOP SETTINGS","#38bdf8"))
        gb = QGroupBox()
        gb.setStyleSheet("QGroupBox{background:#0c0c1e;border:1px solid #1a1a38;"
                         "border-radius:8px;padding:10px;}")
        gl = QVBoxLayout(gb); gl.setSpacing(8)

        def spinrow(lbl, widget, color):
            r = QHBoxLayout()
            l = QLabel(lbl); l.setFont(QFont("Courier New",9))
            l.setStyleSheet("color:#555;background:transparent;border:none;")
            widget.setFixedWidth(90)
            widget.setStyleSheet(f"background:#080818;color:{color};"
                                 f"border:1px solid #1a1a38;border-radius:4px;padding:4px 8px;"
                                 f"font-family:Courier New;font-size:11px;")
            r.addWidget(l); r.addStretch(); r.addWidget(widget); return r

        self.max_loops  = QSpinBox(); self.max_loops.setRange(1,8);  self.max_loops.setValue(3)
        self.stop_score = QDoubleSpinBox()
        self.stop_score.setRange(5.0,10.0); self.stop_score.setSingleStep(0.5)
        self.stop_score.setValue(8.0); self.stop_score.setSuffix("/10"); self.stop_score.setDecimals(1)
        self.debate_chk = QCheckBox("⚔  Debate Mode")
        self.debate_chk.setChecked(True); self.debate_chk.setFont(QFont("Courier New",10))
        self.debate_chk.setStyleSheet("color:#c084fc;background:transparent;")
        gl.addLayout(spinrow("Max Loops:", self.max_loops, "#f59e0b"))
        gl.addLayout(spinrow("Stop Score:", self.stop_score, "#22c55e"))
        gl.addWidget(self.debate_chk)
        lay.addWidget(gb)
        lay.addWidget(sec("◈ EXPORT","#22c55e"))
        self.export_cb = cmb(["Word Document (.docx)","Markdown (.md)","HTML (.html)","Plain Text (.txt)"])
        lay.addWidget(self.export_cb)

        def styled_btn(txt, grad1, grad2, text_color="#000"):
            b = QPushButton(txt); b.setFont(QFont("Courier New",11,QFont.Bold))
            b.setCursor(Qt.PointingHandCursor); b.setFixedHeight(42)
            b.setStyleSheet(f"""QPushButton{{
              background:qlineargradient(x1:0,y1:0,x2:1,y2:0,stop:0 {grad1},stop:1 {grad2});
              color:{text_color};border:none;border-radius:8px;letter-spacing:1px;}}
              QPushButton:hover{{opacity:0.9;}}
              QPushButton:disabled{{background:#111130;color:#333;}}""")
            return b

        self.run_btn = styled_btn("▶  LAUNCH RESEARCH LAB","#c05000","#f59e0b")
        self.run_btn.setFixedHeight(48)
        self.run_btn.clicked.connect(self.start)
        lay.addWidget(self.run_btn)
        self.abort_btn = QPushButton("■  ABORT PIPELINE")
        self.abort_btn.setFont(QFont("Courier New",10,QFont.Bold))
        self.abort_btn.setCursor(Qt.PointingHandCursor); self.abort_btn.setFixedHeight(36)
        self.abort_btn.setEnabled(False)
        self.abort_btn.setStyleSheet("""QPushButton{background:#160404;color:#f87171;
          border:1px solid #f87171;border-radius:7px;}
          QPushButton:hover{background:#200606;}
          QPushButton:disabled{opacity:0.3;}""")
        self.abort_btn.clicked.connect(self.abort)
        lay.addWidget(self.abort_btn)
        self.export_btn = styled_btn("⬇  EXPORT BEST PAPER","#064014","#0a5a1f","#22c55e")
        self.export_btn.setEnabled(False)
        self.export_btn.clicked.connect(self.export)
        lay.addWidget(self.export_btn)
        lay.addStretch()
        self.particles = ParticleWidget(w)
        self.particles.setGeometry(w.rect())
        self.particles.resize(320,900)
        return w

    def _make_center(self):
        w = QWidget(); w.setStyleSheet("background:#06060d;")
        lay = QVBoxLayout(w); lay.setContentsMargins(0,0,0,0); lay.setSpacing(0)
        bar = QFrame(); bar.setFixedHeight(38)
        bar.setStyleSheet("background:#090916;border-bottom:1px solid #12122a;")
        bl = QHBoxLayout(bar); bl.setContentsMargins(16,0,16,0)
        icon_lbl = QLabel("◈  AGENT CONVERSATION FEED  —  LIVE STREAM")
        icon_lbl.setFont(QFont("Courier New",9,QFont.Bold))
        icon_lbl.setStyleSheet("color:#202050;letter-spacing:3px;")
        self.iter_lbl = QLabel("")
        self.iter_lbl.setFont(QFont("Courier New",10,QFont.Bold))
        self.iter_lbl.setStyleSheet("color:#3a3a80;")
        self.wc_total = QLabel("words: 0")
        self.wc_total.setFont(QFont("Courier New",9))
        self.wc_total.setStyleSheet("color:#202040;")
        bl.addWidget(icon_lbl); bl.addStretch()
        bl.addWidget(self.wc_total); bl.addSpacing(12)
        bl.addWidget(self.iter_lbl)
        lay.addWidget(bar)
        self.glow_bar = GlowBar("#f59e0b", 10)
        lay.addWidget(self.glow_bar)
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.scroll.setStyleSheet("""QScrollArea{border:none;background:#06060d;}
          QScrollBar:vertical{background:#09091a;width:7px;border-radius:3px;}
          QScrollBar::handle:vertical{background:#1a1a38;border-radius:3px;min-height:20px;}""")
        self.chat_w = QWidget(); self.chat_w.setStyleSheet("background:#06060d;")
        self.chat_lay = QVBoxLayout(self.chat_w)
        self.chat_lay.setContentsMargins(8,8,8,32); self.chat_lay.setSpacing(3)
        self.chat_lay.addStretch()
        self.scroll.setWidget(self.chat_w)
        lay.addWidget(self.scroll, 1)
        return w

    def _make_right(self):
        w = QWidget(); w.setMaximumWidth(340)
        w.setStyleSheet("background:#07070e;")
        tabs = QTabWidget()
        tabs.setStyleSheet("""
          QTabWidget::pane{border:none;background:#07070e;}
          QTabBar{background:#05050c;}
          QTabBar::tab{background:#09091a;color:#333;font-family:Courier New;
            font-size:9px;padding:6px 12px;border-bottom:2px solid transparent;}
          QTabBar::tab:selected{background:#0c0c1e;color:#f59e0b;
            border-bottom:2px solid #f59e0b;}
        """)
        # Scores Tab
        scores_tab = QWidget(); scores_tab.setStyleSheet("background:#07070e;")
        stl = QVBoxLayout(scores_tab); stl.setContentsMargins(12,12,12,12); stl.setSpacing(10)
        rings_row = QHBoxLayout()
        self.ring_latest = ScoreRing("LATEST","#e879f9",100)
        self.ring_best   = ScoreRing("BEST",  "#22c55e",100)
        rings_row.addWidget(self.ring_latest); rings_row.addStretch()
        rings_row.addWidget(self.ring_best)
        stl.addLayout(rings_row)
        lbl = QLabel("◈ SCORE PROGRESSION")
        lbl.setFont(QFont("Courier New",8,QFont.Bold))
        lbl.setStyleSheet("color:#e879f9;letter-spacing:2px;background:transparent;")
        stl.addWidget(lbl)
        self.score_graph = ScoreGraph()
        stl.addWidget(self.score_graph)
        thresh_row = QHBoxLayout()
        tl = QLabel("THRESHOLD:")
        tl.setFont(QFont("Courier New",9)); tl.setStyleSheet("color:#333;background:transparent;")
        self.thresh_display = QLabel("8.0/10")
        self.thresh_display.setFont(QFont("Courier New",10,QFont.Bold))
        self.thresh_display.setStyleSheet("color:#22c55e;background:transparent;")
        thresh_row.addWidget(tl); thresh_row.addStretch(); thresh_row.addWidget(self.thresh_display)
        stl.addLayout(thresh_row)
        self.score_bar = GlowBar("#22c55e", 10)
        stl.addWidget(self.score_bar)
        stl.addStretch()
        tabs.addTab(scores_tab, "★ SCORES")
        # Pipeline Tab
        pipe_tab = QWidget(); pipe_tab.setStyleSheet("background:#07070e;")
        ptl = QVBoxLayout(pipe_tab); ptl.setContentsMargins(12,12,12,12); ptl.setSpacing(6)
        lbl2 = QLabel("◈ AGENT PIPELINE STATUS")
        lbl2.setFont(QFont("Courier New",8,QFont.Bold))
        lbl2.setStyleSheet("color:#60a5fa;letter-spacing:2px;background:transparent;")
        ptl.addWidget(lbl2)
        self.pipe_labels = {}
        for ag in AGENTS:
            row = QHBoxLayout()
            ic = QLabel(ag["icon"]); ic.setFont(QFont("Segoe UI Emoji",12))
            ic.setFixedWidth(22); ic.setStyleSheet("background:transparent;border:none;")
            nm = QLabel(ag["name"]); nm.setFont(QFont("Courier New",9))
            nm.setStyleSheet(f"color:{ag['color']}40;background:transparent;border:none;")
            st = QLabel("○"); st.setFont(QFont("Courier New",12))
            st.setAlignment(Qt.AlignRight)
            st.setStyleSheet("color:#1a1a30;background:transparent;border:none;")
            pb = QProgressBar(); pb.setFixedHeight(3); pb.setTextVisible(False)
            pb.setRange(0,0); pb.setVisible(False)
            pb.setStyleSheet(f"QProgressBar{{background:#111;border:none;border-radius:1px;}}"
                             f"QProgressBar::chunk{{background:{ag['color']};border-radius:1px;}}")
            row.addWidget(ic); row.addWidget(nm); row.addStretch()
            row.addWidget(pb,2); row.addWidget(st)
            ptl.addLayout(row)
            self.pipe_labels[ag["id"]] = (nm, st, pb)
        ptl.addStretch()
        info_lbl_h = QLabel("◈ SESSION")
        info_lbl_h.setFont(QFont("Courier New",8,QFont.Bold))
        info_lbl_h.setStyleSheet("color:#fbbf24;letter-spacing:2px;background:transparent;")
        ptl.addWidget(info_lbl_h)
        self.info_lbl = QLabel("Status: IDLE\nIterations: 0\nBest Score: —")
        self.info_lbl.setFont(QFont("Courier New",9))
        self.info_lbl.setStyleSheet("color:#333;background:transparent;line-height:170%;")
        self.info_lbl.setWordWrap(True)
        ptl.addWidget(self.info_lbl)
        tabs.addTab(pipe_tab, "⚙ PIPELINE")
        # Knowledge Graph Tab
        kg_tab = QWidget(); kg_tab.setStyleSheet("background:#07070e;")
        kgl = QVBoxLayout(kg_tab); kgl.setContentsMargins(8,8,8,8)
        lbl3 = QLabel("◈ KNOWLEDGE GRAPH")
        lbl3.setFont(QFont("Courier New",8,QFont.Bold))
        lbl3.setStyleSheet("color:#34d399;letter-spacing:2px;background:transparent;")
        kgl.addWidget(lbl3)
        self.kg_widget = KnowledgeGraph()
        kgl.addWidget(self.kg_widget, 1)
        tabs.addTab(kg_tab, "🕸 GRAPH")
        # Memory Tab
        mem_tab = QWidget(); mem_tab.setStyleSheet("background:#07070e;")
        meml = QVBoxLayout(mem_tab); meml.setContentsMargins(10,10,10,10); meml.setSpacing(6)
        lbl4 = QLabel("◈ AGENT MEMORY")
        lbl4.setFont(QFont("Courier New",8,QFont.Bold))
        lbl4.setStyleSheet("color:#a78bfa;letter-spacing:2px;background:transparent;")
        meml.addWidget(lbl4)
        self.memory_cb = QComboBox(); self.memory_cb.setFont(QFont("Courier New",9))
        self.memory_cb.setStyleSheet("""QComboBox{background:#0c0c1e;color:#c0cce0;
          border:1px solid #1a1a38;border-radius:4px;padding:5px 8px;}
          QComboBox::drop-down{border:none;}
          QComboBox QAbstractItemView{background:#0c0c1e;color:#c0cce0;
            border:1px solid #1a1a38;selection-background-color:#a78bfa30;}""")
        self.memory_view = QTextBrowser()
        self.memory_view.setFont(QFont("Courier New",10))
        self.memory_view.setStyleSheet("""QTextBrowser{background:#05050e;color:#6a7a9a;
          border:1px solid #111130;border-radius:4px;}
          QScrollBar:vertical{background:#09091a;width:5px;border-radius:2px;}
          QScrollBar::handle:vertical{background:#1a1a38;border-radius:2px;}""")
        meml.addWidget(self.memory_cb)
        meml.addWidget(self.memory_view, 1)
        self.memory_cb.currentTextChanged.connect(self._update_memory_view)
        tabs.addTab(mem_tab, "🧠 MEMORY")
        lay = QVBoxLayout(w); lay.setContentsMargins(0,0,0,0)
        lay.addWidget(tabs)
        return w

    def _make_statusbar(self):
        b = QFrame(); b.setFixedHeight(30)
        b.setStyleSheet("background:#04040c;border-top:1px solid #0e0e28;")
        lay = QHBoxLayout(b); lay.setContentsMargins(16,0,16,0)
        self.status_lbl = QLabel("● IDLE  —  Configure API key and topic, then launch")
        self.status_lbl.setFont(QFont("Courier New",9))
        self.status_lbl.setStyleSheet("color:#2a2a50;")
        lay.addWidget(self.status_lbl); lay.addStretch()
        agent_count = QLabel(f"Life Of Research  ·  By Pranay Mahendrakar  ·  AGENTS: {len(AGENTS)}  ·  14-STAGE AUTONOMOUS PIPELINE")
        agent_count.setFont(QFont("Courier New",8))
        agent_count.setStyleSheet("color:#181838;")
        lay.addWidget(agent_count)
        return b

    # ═══════════════════════════════════════════════════════════════
    #   PIPELINE ORCHESTRATION
    # ═══════════════════════════════════════════════════════════════

    def start(self):
        api   = self.api_inp.text().strip()
        topic = self.topic_edit.toPlainText().strip()
        if not api:   return self._sset("⚠ Enter your OpenAI API key","#f87171")
        if not topic: return self._sset("⚠ Enter a research topic","#f87171")
        self.running     = True; self.abort_flag = False
        self.iteration   = 0;    self.scores = []
        self.memory      = {"topic":topic,
                             "paper_type":self.type_cb.currentText(),
                             "citation_style":self.cite_cb.currentText()}
        self.outputs     = {}; self.bubbles = {}
        self.best_paper  = ""; self.best_score = 0.0; self.best_iter = 0
        self._ach_shown  = set(); self._elapsed = 0; self._total_words = 0
        self.score_graph.clear()
        self.ring_latest.set_value(0); self.ring_best.set_value(0)
        self.thresh_display.setText(f"{self.stop_score.value():.1f}/10")
        self._clear_chat(); self._reset_pipe()
        self.glow_bar.set_value(0)
        self.score_bar.set_value(0)
        self.memory_cb.clear()
        self.run_btn.setEnabled(False); self.abort_btn.setEnabled(True)
        self.export_btn.setEnabled(False)
        self.particles.start()
        self.kg_widget.set_topic(topic)
        self._run("director", [
            {"role":"system","content":AGENTS_MAP["director"]["system"]},
            {"role":"user","content":
             f'Research Topic: "{topic}"\nPaper Type: {self.memory["paper_type"]}\n'
             f'Citation Style: {self.memory["citation_style"]}\nCreate research direction plan.'}
        ], self._cb_director)

    def _run(self, agent_id, messages, callback):
        if self.abort_flag: return self._finish(True)
        self._set_pipe(agent_id,"running")
        bubble = self._add_bubble(agent_id, self.iteration if self.iteration>0 else None)
        self._sset(f"◈  {AGENTS_MAP[agent_id]['name']}  —  {AGENTS_MAP[agent_id]['role']}...",
                  AGENTS_MAP[agent_id]["color"])
        w = AgentWorker(agent_id, messages,
                        self.api_inp.text().strip(), self.model_cb.currentText())
        w.token_received.connect(lambda aid,tok: (
            bubble.append(tok),
            self.wc_total.setText(f"words generated: {self._count_total_words():,}")
        ))
        w.agent_done.connect(callback)
        w.agent_error.connect(self._on_error)
        if agent_id == "evaluator":
            w.score_found.connect(self._on_score)
        w.start(); self.worker = w

    def _count_total_words(self):
        return sum(len(b._raw.split()) for b in self.bubbles.values())

    def _cb_director(self, aid, text):
        self.memory["director"] = text; self._done(aid)
        self._update_memory_dropdown()
        self._run("literature", [
            {"role":"system","content":AGENTS_MAP["literature"]["system"]},
            {"role":"user","content":
             f'Topic: "{self.memory["topic"]}"\nPlan:\n{text}\n'
             f'Generate literature analysis with {self.memory["citation_style"]} citations.'}
        ], self._cb_literature)

    def _cb_literature(self, aid, text):
        self.memory["literature"] = text; self._done(aid)
        self._update_memory_dropdown()
        self._run("hypothesis", [
            {"role":"system","content":AGENTS_MAP["hypothesis"]["system"]},
            {"role":"user","content":
             f'Topic: "{self.memory["topic"]}"\nPlan:\n{self.memory["director"]}\n'
             f'Literature:\n{text}\nGenerate hypotheses and experiment design.'}
        ], self._cb_hypothesis)

    def _cb_hypothesis(self, aid, text):
        self.memory["hypothesis"] = text; self._done(aid)
        self._update_memory_dropdown()
        self._start_loop()

    def _start_loop(self):
        if self.abort_flag: return self._finish(True)
        self.iteration += 1
        self.outputs[self.iteration] = {}
        max_l = self.max_loops.value()
        self.iter_lbl.setText(f"ITERATION  {self.iteration}  /  {max_l}")
        self._add_div(f"⟳  LOOP ITERATION  {self.iteration}  ·  TARGET {self.stop_score.value():.1f}/10")
        self.glow_bar.set_value(int(self.iteration/max_l*100))
        self._reset_loop_pipe()
        prev = ""
        if self.iteration > 1:
            prev = (
                f"\n\n---\n### Previous Reviewer Feedback (MUST address ALL points):\n\n"
                f"**Methods Review:**\n{self.memory.get('reviewer_methods','')}\n\n"
                f"**Novelty Review:**\n{self.memory.get('reviewer_novelty','')}\n\n"
                f"**Harsh Review:**\n{self.memory.get('reviewer_harsh','')}\n\n"
                f"**Rebuttal/Conceded:**\n{self.memory.get('defender','')}"
            )
        self._run("writer", [
            {"role":"system","content":AGENTS_MAP["writer"]["system"]},
            {"role":"user","content":
             f'Topic: "{self.memory["topic"]}"\nType: {self.memory["paper_type"]}\n'
             f'Citation: {self.memory["citation_style"]}\n\n'
             f'Research Plan:\n{self.memory.get("director","")}\n\n'
             f'Literature:\n{self.memory.get("literature","")}\n\n'
             f'Hypotheses:\n{self.memory.get("hypothesis","")}'
             + prev}
        ], self._cb_writer)

    def _cb_writer(self, aid, text):
        self.memory["writer_current"] = text
        self.outputs[self.iteration]["writer"] = text; self._done(aid)
        self._run("reviewer_methods", [
            {"role":"system","content":AGENTS_MAP["reviewer_methods"]["system"]},
            {"role":"user","content":f"Review this paper:\n\n{text}"}
        ], self._cb_rev_meth)

    def _cb_rev_meth(self, aid, text):
        self.memory["reviewer_methods"] = text
        self.outputs[self.iteration]["reviewer_methods"] = text; self._done(aid)
        self._run("reviewer_novelty", [
            {"role":"system","content":AGENTS_MAP["reviewer_novelty"]["system"]},
            {"role":"user","content":f"Review this paper:\n\n{self.memory['writer_current']}"}
        ], self._cb_rev_nov)

    def _cb_rev_nov(self, aid, text):
        self.memory["reviewer_novelty"] = text
        self.outputs[self.iteration]["reviewer_novelty"] = text; self._done(aid)
        self._run("reviewer_harsh", [
            {"role":"system","content":AGENTS_MAP["reviewer_harsh"]["system"]},
            {"role":"user","content":f"Review this paper:\n\n{self.memory['writer_current']}"}
        ], self._cb_rev_harsh)

    def _cb_rev_harsh(self, aid, text):
        self.memory["reviewer_harsh"] = text
        self.outputs[self.iteration]["reviewer_harsh"] = text; self._done(aid)
        if self.debate_chk.isChecked():
            self._run("defender", [
                {"role":"system","content":AGENTS_MAP["defender"]["system"]},
                {"role":"user","content":
                 f'Paper:\n{self.memory["writer_current"]}\n\n'
                 f'Methods Review:\n{self.memory["reviewer_methods"]}\n\n'
                 f'Novelty Review:\n{self.memory["reviewer_novelty"]}\n\n'
                 f'Harsh Review:\n{text}'}
            ], self._cb_defender)
        else:
            self._run_corrector()

    def _cb_defender(self, aid, text):
        self.memory["defender"] = text
        self.outputs[self.iteration]["defender"] = text; self._done(aid)
        self._run_corrector()

    def _run_corrector(self):
        self._run("corrector", [
            {"role":"system","content":AGENTS_MAP["corrector"]["system"]},
            {"role":"user","content":
             f'Draft:\n{self.memory["writer_current"]}\n\n'
             f'Methods Review:\n{self.memory.get("reviewer_methods","")}\n\n'
             f'Novelty Review:\n{self.memory.get("reviewer_novelty","")}\n\n'
             f'Harsh Review:\n{self.memory.get("reviewer_harsh","")}\n\n'
             f'Rebuttal:\n{self.memory.get("defender","Skipped.")}'}
        ], self._cb_corrector)

    def _cb_corrector(self, aid, text):
        self.memory["corrector"] = text
        self.outputs[self.iteration]["corrector"] = text; self._done(aid)
        self._run("humanizer", [
            {"role":"system","content":AGENTS_MAP["humanizer"]["system"]},
            {"role":"user","content":f"Humanize this paper:\n\n{text}"}
        ], self._cb_humanizer)

    def _cb_humanizer(self, aid, text):
        self.memory["humanizer"] = text
        self.outputs[self.iteration]["humanizer"] = text; self._done(aid)
        self._run("evaluator", [
            {"role":"system","content":AGENTS_MAP["evaluator"]["system"]},
            {"role":"user","content":
             f'Final paper:\n{text}\n\n'
             f'Methods review:\n{self.memory.get("reviewer_methods","")}\n\n'
             f'Novelty review:\n{self.memory.get("reviewer_novelty","")}\n\n'
             f'Harsh review:\n{self.memory.get("reviewer_harsh","")}\n\n'
             f'Give final publication score.'}
        ], self._cb_evaluator)

    def _on_score(self, score):
        self.scores.append(score)
        self.score_graph.add_score(score)
        self.ring_latest.set_value(score)
        if score > self.best_score:
            self.best_score = score; self.best_iter = self.iteration
            self.best_paper = self.memory.get("humanizer","")
            self.ring_best.set_value(score)
        self.score_bar.set_value(min(int(score/10*100),100))
        for threshold,icon,msg,color in ACHIEVEMENTS:
            if score >= threshold and threshold not in self._ach_shown:
                self._ach_shown.add(threshold)
                self.achievement.show_achievement(icon, f"ACHIEVEMENT UNLOCKED: {msg}", color)
        self._update_info()

    def _cb_evaluator(self, aid, text):
        self.memory["evaluator"] = text
        self.outputs[self.iteration]["evaluator"] = text; self._done(aid)
        self._update_memory_dropdown()
        score     = self.scores[-1] if self.scores else 0
        threshold = self.stop_score.value()
        max_l     = self.max_loops.value()
        if self.abort_flag:
            self._finish(True)
        elif score >= threshold:
            self._sset(f"✓ COMPLETE — Score {score:.1f} reached {threshold:.1f} threshold!","#22c55e")
            self.achievement.show_achievement("🎯","LIFE OF RESEARCH — Target Score Reached!","#22c55e")
            self._finish(False)
        elif self.iteration >= max_l:
            self._sset(f"✓ MAX LOOPS — Best: {self.best_score:.1f}/10 at iteration {self.best_iter}","#f59e0b")
            self._finish(False)
        else:
            self._sset(f"⟳ Score {score:.1f} < {threshold:.1f} — launching iteration {self.iteration+1}...","#60a5fa")
            self._reset_loop_pipe()
            QTimer.singleShot(700, self._start_loop)

    def _on_error(self, aid, msg):
        if aid in self.bubbles: self.bubbles[aid].error(msg)
        self._set_pipe(aid,"error")
        self._sset(f"✗ ERROR in [{AGENTS_MAP[aid]['name']}]: {msg[:120]}","#f87171")
        self._finish(True)

    def _finish(self, aborted):
        self.running = False
        self.particles.stop()
        self.run_btn.setEnabled(True); self.abort_btn.setEnabled(False)
        self.export_btn.setEnabled(not aborted and bool(self.best_paper))
        if not aborted: self.glow_bar.set_value(100)
        self._update_info()

    def abort(self):
        self.abort_flag = True
        if self.worker: self.worker.stop()
        self._sset("■ Aborting pipeline...","#f87171")

    # ═══════════════════════════════════════════════════════════════
    #   EXPORT
    # ═══════════════════════════════════════════════════════════════

    def export(self):
        topic = self.topic_edit.toPlainText().strip()
        ts    = datetime.now().strftime("%Y-%m-%d_%H-%M")
        paper = self.best_paper
        bi    = self.best_iter; bs = self.best_score
        it    = self.outputs.get(bi, {})
        reviews = {rid: it.get(rid,"") for rid in REVIEWER_IDS}
        fmt   = self.export_cb.currentIndex()
        exts  = [".docx",".md",".html",".txt"]
        mimes = ["Word Document (*.docx)","Markdown (*.md)","HTML (*.html)","Text (*.txt)"]
        fname,_ = QFileDialog.getSaveFileName(self,"Export Research Paper",
                                               f"research_{ts}{exts[fmt]}", mimes[fmt])
        if not fname: return
        if fmt == 0:  # DOCX
            try:
                export_docx(fname, topic, paper,
                            self.memory.get("director",""),
                            self.memory.get("literature",""),
                            self.memory.get("hypothesis",""),
                            reviews,
                            it.get("evaluator",""),
                            bi, bs, ts)
                self._sset(f"✓ DOCX exported → {os.path.basename(fname)}","#22c55e")
                self.achievement.show_achievement("📄","Paper exported as Word Document!","#60a5fa")
            except Exception as e:
                self._sset(f"✗ DOCX export failed: {str(e)[:100]}","#f87171")
        elif fmt == 1:  # Markdown
            content = (f"# Research Paper: {topic}\n\n"
                      f"*Life Of Research  ·  By Pranay Mahendrakar  ·  {ts}  ·  Score: {bs:.1f}/10*\n\n---\n\n"
                      f"## Research Direction\n\n{self.memory.get('director','')}\n\n---\n\n"
                      f"## Literature Analysis\n\n{self.memory.get('literature','')}\n\n---\n\n"
                      f"## Hypotheses & Experiments\n\n{self.memory.get('hypothesis','')}\n\n---\n\n"
                      f"## Final Paper — Iteration {bi}\n\n{paper}\n\n---\n\n")
            for rid in REVIEWER_IDS:
                content += f"## {AGENTS_MAP[rid]['name']}\n\n{reviews.get(rid,'')}\n\n---\n\n"
            content += f"## Final Evaluation\n\n{it.get('evaluator','')}\n"
            with open(fname,"w",encoding="utf-8") as f: f.write(content)
            self._sset(f"✓ Markdown exported → {os.path.basename(fname)}","#22c55e")
        elif fmt == 2:  # HTML
            def e(s): return s.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            sections = ""
            for title, content in [
                ("Research Direction", self.memory.get("director","")),
                ("Literature Analysis", self.memory.get("literature","")),
                ("Hypotheses & Experiments", self.memory.get("hypothesis","")),
                (f"Final Paper — Iteration {bi}", paper),
                ("Methods Review", reviews.get("reviewer_methods","")),
                ("Novelty Review", reviews.get("reviewer_novelty","")),
                ("Harsh Review",   reviews.get("reviewer_harsh","")),
                ("Final Evaluation", it.get("evaluator","")),
            ]:
                sections += f'<h2>{e(title)}</h2><pre>{e(content)}</pre>\n'
            content = f"""<!DOCTYPE html><html lang="en"><head>
<meta charset="UTF-8"><title>{e(topic)}</title>
<style>
body{{font-family:'Segoe UI',Arial,serif;max-width:1000px;margin:40px auto;
     padding:24px;line-height:1.9;color:#1a1a2e;background:#fffdf7}}
h1{{text-align:center;font-size:28px;color:#1a237e}}
h2{{font-size:18px;color:#c05000;border-bottom:2px solid #c05000;
    padding-bottom:6px;margin-top:40px}}
.meta{{text-align:center;color:#888;margin-bottom:40px;font-size:14px}}
.score{{background:#c05000;color:#fff;padding:3px 14px;border-radius:20px;
        font-weight:bold;font-size:14px}}
pre{{background:#f5f5fb;padding:18px;border-radius:8px;font-size:12.5px;
     white-space:pre-wrap;word-break:break-word;
     border-left:4px solid #c05000;font-family:'Courier New',monospace}}
</style></head><body>
<h1>{e(topic)}</h1>
<div class="meta">Life Of Research · By Pranay Mahendrakar · {ts} · <span class="score">Score {bs:.1f}/10</span></div>
{sections}
</body></html>"""
            with open(fname,"w",encoding="utf-8") as f: f.write(content)
            self._sset(f"✓ HTML exported → {os.path.basename(fname)}","#22c55e")
        else:  # txt
            sep = "═"*80+"\n"
            content = (f"RESEARCH PAPER: {topic}\n{sep}"
                      f"Generated: {ts}  |  Score: {bs:.1f}/10  |  Iter {bi}\n{sep}\n"
                      f"RESEARCH DIRECTION\n{sep}{self.memory.get('director','')}\n\n"
                      f"FINAL PAPER\n{sep}{paper}\n\n")
            with open(fname,"w",encoding="utf-8") as f: f.write(content)
            self._sset(f"✓ Text exported → {os.path.basename(fname)}","#22c55e")

    # ═══════════════════════════════════════════════════════════════
    #   UI HELPERS
    # ═══════════════════════════════════════════════════════════════

    def _add_bubble(self, agent_id, iteration=None):
        b = ChatBubble(agent_id, iteration)
        self.bubbles[agent_id] = b
        self.chat_lay.insertWidget(self.chat_lay.count()-1, b)
        QTimer.singleShot(80, lambda:
            self.scroll.verticalScrollBar().setValue(
            self.scroll.verticalScrollBar().maximum()))
        return b

    def _add_div(self, label):
        self.chat_lay.insertWidget(self.chat_lay.count()-1, IterDiv(label))

    def _clear_chat(self):
        self.bubbles = {}
        while self.chat_lay.count() > 1:
            item = self.chat_lay.takeAt(0)
            if item.widget(): item.widget().deleteLater()

    def _done(self, aid):
        if aid in self.bubbles: self.bubbles[aid].done()
        self._set_pipe(aid,"done")
        self._update_memory_dropdown()

    def _reset_pipe(self):
        for aid,(nm,st,pb) in self.pipe_labels.items():
            ag = AGENTS_MAP[aid]
            nm.setStyleSheet(f"color:{ag['color']}30;background:transparent;border:none;")
            st.setText("○"); st.setStyleSheet("color:#141430;background:transparent;border:none;")
            pb.setVisible(False)

    def _reset_loop_pipe(self):
        loop_ids = (["writer"]+REVIEWER_IDS+["defender","corrector","humanizer","evaluator"])
        for aid in loop_ids:
            if aid in self.pipe_labels:
                ag = AGENTS_MAP[aid]; nm,st,pb = self.pipe_labels[aid]
                nm.setStyleSheet(f"color:{ag['color']}30;background:transparent;border:none;")
                st.setText("○"); st.setStyleSheet("color:#141430;background:transparent;border:none;")
                pb.setVisible(False)

    def _set_pipe(self, aid, state):
        if aid not in self.pipe_labels: return
        ag = AGENTS_MAP[aid]; nm,st,pb = self.pipe_labels[aid]
        if state=="running":
            nm.setStyleSheet(f"color:{ag['color']};background:transparent;border:none;font-weight:bold;")
            st.setText("⟳"); st.setStyleSheet(f"color:{ag['color']};background:transparent;border:none;")
            pb.setVisible(True)
        elif state=="done":
            nm.setStyleSheet(f"color:{ag['color']}80;background:transparent;border:none;")
            st.setText("✓"); st.setStyleSheet("color:#22c55e;background:transparent;border:none;")
            pb.setVisible(False)
        elif state=="error":
            st.setText("✗"); st.setStyleSheet("color:#f87171;background:transparent;border:none;")
            pb.setVisible(False)

    def _update_info(self):
        best_i = f"Iteration {self.best_iter}" if self.best_iter else "—"
        wc = self._count_total_words()
        self.wc_total.setText(f"words generated: {wc:,}")
        self.info_lbl.setText(
            f"Status: {'🔥 RUNNING' if self.running else '✓ DONE'}\n"
            f"Iterations: {self.iteration}\n"
            f"Best Score: {self.best_score:.1f}/10\n"
            f"Best Iter: {best_i}\n"
            f"Total Words: {wc:,}\n"
            f"Threshold: {self.stop_score.value():.1f}/10"
        )
        self.info_lbl.setStyleSheet("color:#4a4a80;background:transparent;line-height:170%;")

    def _update_memory_dropdown(self):
        keys = [k for k in self.memory.keys() if isinstance(self.memory[k],str) and len(self.memory[k])>20]
        current = self.memory_cb.currentText()
        self.memory_cb.blockSignals(True)
        self.memory_cb.clear()
        self.memory_cb.addItems(keys)
        if current in keys: self.memory_cb.setCurrentText(current)
        self.memory_cb.blockSignals(False)
        self._update_memory_view()

    def _update_memory_view(self):
        key = self.memory_cb.currentText()
        if key and key in self.memory:
            text = self.memory[key]
            html = md_to_html(str(text)[:4000], "#a78bfa")
            self.memory_view.setHtml(
                f'<html><body style="background:#05050e;padding:8px;">{html}</body></html>')

    def _sset(self, msg, color="#c0cce0"):
        self.status_lbl.setText(msg)
        self.status_lbl.setStyleSheet(f"color:{color};font-family:Courier New;font-size:9px;")


# ═══════════════════════════════════════════════════════════════════
#   ENTRY POINT
# ═══════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    app.setApplicationName("Life Of Research — By Pranay Mahendrakar")
    w = LifeOfResearchWindow()
    w.show()
    sys.exit(app.exec())
